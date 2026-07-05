#!/usr/bin/env python3
from __future__ import annotations

import base64
import json
import re
import shutil
import tempfile
import threading
import wave
import webbrowser
from dataclasses import dataclass
from email.message import EmailMessage
from pathlib import Path
from urllib.parse import quote
from tkinter import filedialog, messagebox, simpledialog, ttk
import tkinter as tk

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import fitz
import numpy as np
import sounddevice as sd
from faster_whisper import WhisperModel
from google import genai
from google.genai import types
from pypdf import PdfReader, PdfWriter
from pypdf.generic import BooleanObject, NameObject


DEFAULT_GEMINI_MODEL = "gemini-flash-latest"
GMAIL_SCOPES = ["https://www.googleapis.com/auth/gmail.compose"]
EDIT_OPERATIONS = {
    "insert_after_sentence",
    "insert_before_sentence",
    "replace_sentence",
    "replace_paragraph",
    "insert_paragraph_after",
    "insert_paragraph_before",
    "clarify",
}


@dataclass
class DocxEditPlan:
    operation: str
    paragraph_index: int
    anchor_text: str
    new_text: str
    explanation: str


class GeminiFormFillerApp:
    def __init__(self, root: tk.Tk) -> None:
        self.root = root
        self.root.title("Whisper Voice To Form - V6")
        self.root.geometry("1280x820")

        self.app_dir = Path(__file__).resolve().parent
        self.api_key_file = self.app_dir / ".gemini-api-key"
        self.gmail_credentials_file = self.app_dir / "gmail-credentials.json"
        self.gmail_token_file = self.app_dir / ".gmail-token.json"
        self.model_cache_dir = self.app_dir / "model-cache"
        self.templates_dir = self.app_dir / "templates"
        self.model_cache_dir.mkdir(exist_ok=True)
        self.templates_dir.mkdir(exist_ok=True)

        self.pdf_path: Path | None = None
        self.pdf_document: fitz.Document | None = None
        self.docx_path: Path | None = None
        self.docx_document: Document | None = None
        self.selected_docx_paragraph_index: int | None = None
        self.pdf_fields: dict[str, dict] = {}
        self.pdf_field_locations: dict[str, list[tuple[int, fitz.Rect]]] = {}
        self.pdf_widget_values: dict[str, str] = {}
        self.field_values: dict[str, str] = {}
        self.current_page_index = 0
        self.preview_zoom = 1.25
        self.preview_image: tk.PhotoImage | None = None
        self.preview_editor: ttk.Entry | None = None
        self.preview_editor_window: int | None = None
        self.preview_editor_field: str | None = None
        self.recording = False
        self.record_stream = None
        self.recorded_chunks: list[np.ndarray] = []
        self.sample_rate = 16000
        self.whisper_model: WhisperModel | None = None
        self.is_busy = False

        self.api_key_var = tk.StringVar(value=self._load_api_key())
        self.gemini_model_var = tk.StringVar(value=DEFAULT_GEMINI_MODEL)
        self.whisper_model_var = tk.StringVar(value="base")
        self.pdf_status_var = tk.StringVar(value="No PDF loaded")
        self.docx_status_var = tk.StringVar(value="No Word document loaded")
        self.page_status_var = tk.StringVar(value="No page loaded")
        self.recording_status_var = tk.StringVar(value="Mic idle")
        self.status_var = tk.StringVar(value="Ready")
        self.docx_selected_status_var = tk.StringVar(value="No Word paragraph selected")
        self.template_var = tk.StringVar(value="")
        self.email_to_var = tk.StringVar(value="")
        self.email_cc_var = tk.StringVar(value="")
        self.email_bcc_var = tk.StringVar(value="")
        self.email_subject_var = tk.StringVar(value="")
        self.collapsed_panels: set[str] = set()
        self.panel_button_vars = {
            "email": tk.StringVar(value="[v] Email"),
            "fields": tk.StringVar(value="[v] PDF Fields"),
            "docx": tk.StringVar(value="[v] DOCX Edits"),
        }

        self._build_ui()
        self.refresh_template_choices()

    def _build_ui(self) -> None:
        main = ttk.Frame(self.root, padding=12)
        main.pack(fill=tk.BOTH, expand=True)

        settings = ttk.LabelFrame(main, text="AI Settings", padding=10)
        settings.pack(fill=tk.X)

        ttk.Label(settings, text="Google AI Studio API key").grid(row=0, column=0, sticky="w")
        api_entry = ttk.Entry(settings, textvariable=self.api_key_var, show="*", width=58)
        api_entry.grid(row=0, column=1, sticky="ew", padx=(8, 6))
        ttk.Button(settings, text="Save Key", command=self.save_api_key).grid(row=0, column=2)
        ttk.Button(settings, text="Test Key", command=self.test_api_key).grid(row=0, column=3, padx=(6, 0))
        ttk.Button(settings, text="API Key Help", command=self.show_api_key_help).grid(row=0, column=4, padx=(6, 0))

        ttk.Label(settings, text="Gemini model").grid(row=1, column=0, sticky="w", pady=(8, 0))
        ttk.Entry(settings, textvariable=self.gemini_model_var, width=28).grid(row=1, column=1, sticky="w", padx=(8, 0), pady=(8, 0))
        ttk.Label(settings, text="Whisper model").grid(row=1, column=1, sticky="e", padx=(0, 130), pady=(8, 0))
        ttk.Combobox(
            settings,
            textvariable=self.whisper_model_var,
            values=["tiny", "base", "small", "medium", "large-v3"],
            width=12,
            state="readonly",
        ).grid(row=1, column=2, sticky="w", pady=(8, 0))
        settings.columnconfigure(1, weight=1)

        actions = ttk.Frame(main)
        actions.pack(fill=tk.X, pady=(10, 8))

        ttk.Button(actions, text="Upload Fillable PDF", command=self.load_pdf).pack(side=tk.LEFT)
        ttk.Button(actions, text="Auto-Prepare Flat PDF", command=self.auto_prepare_flat_pdf).pack(side=tk.LEFT, padx=(8, 0))
        ttk.Button(actions, text="Review Fields With Gemini", command=self.review_fields_with_gemini).pack(side=tk.LEFT, padx=(8, 0))
        ttk.Button(actions, text="Save As Template", command=self.save_current_as_template).pack(side=tk.LEFT, padx=(8, 0))
        ttk.Button(actions, text="Upload Word Document", command=self.load_docx).pack(side=tk.LEFT, padx=(16, 0))
        ttk.Button(actions, text="Reload Word", command=self.reload_docx).pack(side=tk.LEFT, padx=(8, 0))
        ttk.Button(actions, text="Save Edited DOCX", command=self.save_docx_copy).pack(side=tk.LEFT, padx=(8, 0))
        ttk.Label(actions, text="Template:").pack(side=tk.LEFT, padx=(12, 4))
        self.template_combo = ttk.Combobox(actions, textvariable=self.template_var, width=24, state="readonly")
        self.template_combo.pack(side=tk.LEFT)
        ttk.Button(actions, text="Load Template", command=self.load_selected_template).pack(side=tk.LEFT, padx=(6, 0))

        recording_actions = ttk.Frame(main)
        recording_actions.pack(fill=tk.X, pady=(0, 8))

        ttk.Button(recording_actions, text="Start Recording", command=self.start_recording).pack(side=tk.LEFT)
        ttk.Button(recording_actions, text="Stop and Transcribe", command=self.stop_recording).pack(side=tk.LEFT, padx=(8, 0))
        ttk.Button(recording_actions, text="Fix Grammar / Whisper Text", command=lambda: self.rewrite_transcript_with_gemini("grammar")).pack(side=tk.LEFT, padx=(16, 0))
        ttk.Button(recording_actions, text="Make Email", command=lambda: self.rewrite_transcript_with_gemini("email")).pack(side=tk.LEFT, padx=(8, 0))
        ttk.Button(recording_actions, text="Send Transcript To PDF", command=self.send_transcript_to_ai).pack(side=tk.LEFT, padx=(16, 0))
        ttk.Button(recording_actions, text="Apply Transcript To DOCX", command=self.apply_transcript_to_docx).pack(side=tk.LEFT, padx=(8, 0))
        ttk.Button(recording_actions, text="Export Filled PDF", command=self.export_filled_pdf).pack(side=tk.LEFT, padx=(8, 0))

        ttk.Label(main, textvariable=self.pdf_status_var).pack(anchor="w")
        ttk.Label(main, textvariable=self.docx_status_var).pack(anchor="w", pady=(2, 0))
        ttk.Label(main, textvariable=self.docx_selected_status_var).pack(anchor="w", pady=(2, 0))
        ttk.Label(main, textvariable=self.recording_status_var).pack(anchor="w", pady=(2, 8))

        workspace = ttk.PanedWindow(main, orient=tk.HORIZONTAL)
        workspace.pack(fill=tk.BOTH, expand=True, pady=(0, 8))

        left_column = ttk.PanedWindow(workspace, orient=tk.VERTICAL)
        self.left_column = left_column
        preview = ttk.LabelFrame(workspace, text="PDF Preview", padding=8)
        workspace.add(left_column, weight=1)
        workspace.add(preview, weight=2)

        transcript_panel = ttk.Frame(left_column, padding=(0, 0, 8, 4))
        email_panel = ttk.LabelFrame(left_column, text="Email Draft", padding=(0, 4, 8, 4))
        fields_panel = ttk.Frame(left_column, padding=(0, 4, 8, 0))
        docx_panel = ttk.Frame(left_column, padding=(0, 4, 8, 0))
        left_column.add(transcript_panel, weight=1)
        left_column.add(email_panel, weight=1)
        left_column.add(fields_panel, weight=1)
        left_column.add(docx_panel, weight=2)

        transcript_header = ttk.Frame(transcript_panel)
        transcript_header.pack(fill=tk.X)
        ttk.Label(transcript_header, text="Transcript").pack(side=tk.LEFT)

        panel_toggles = ttk.Frame(transcript_panel)
        panel_toggles.pack(fill=tk.X, pady=(6, 0))
        ttk.Button(panel_toggles, textvariable=self.panel_button_vars["email"], command=lambda: self.toggle_panel("email")).pack(side=tk.LEFT)
        ttk.Button(panel_toggles, textvariable=self.panel_button_vars["fields"], command=lambda: self.toggle_panel("fields")).pack(side=tk.LEFT, padx=(6, 0))
        ttk.Button(panel_toggles, textvariable=self.panel_button_vars["docx"], command=lambda: self.toggle_panel("docx")).pack(side=tk.LEFT, padx=(6, 0))

        self.transcript_text = tk.Text(transcript_panel, wrap=tk.WORD, height=12)
        self.transcript_text.pack(fill=tk.BOTH, expand=True, pady=(6, 0))

        self.collapsible_panes = {
            "email": (email_panel, 1),
            "fields": (fields_panel, 1),
            "docx": (docx_panel, 2),
        }

        email_fields = ttk.Frame(email_panel)
        email_fields.pack(fill=tk.X)
        ttk.Label(email_fields, text="To").grid(row=0, column=0, sticky="w")
        ttk.Entry(email_fields, textvariable=self.email_to_var).grid(row=0, column=1, sticky="ew", padx=(6, 0))
        ttk.Label(email_fields, text="CC").grid(row=1, column=0, sticky="w", pady=(4, 0))
        ttk.Entry(email_fields, textvariable=self.email_cc_var).grid(row=1, column=1, sticky="ew", padx=(6, 0), pady=(4, 0))
        ttk.Label(email_fields, text="BCC").grid(row=2, column=0, sticky="w", pady=(4, 0))
        ttk.Entry(email_fields, textvariable=self.email_bcc_var).grid(row=2, column=1, sticky="ew", padx=(6, 0), pady=(4, 0))
        ttk.Label(email_fields, text="Subject").grid(row=3, column=0, sticky="w", pady=(4, 0))
        ttk.Entry(email_fields, textvariable=self.email_subject_var).grid(row=3, column=1, sticky="ew", padx=(6, 0), pady=(4, 0))
        email_fields.columnconfigure(1, weight=1)

        self.email_body_text = tk.Text(email_panel, wrap=tk.WORD, height=6)
        self.email_body_text.pack(fill=tk.BOTH, expand=True, pady=(6, 4))
        email_buttons = ttk.Frame(email_panel)
        email_buttons.pack(fill=tk.X)
        ttk.Button(email_buttons, text="Connect Gmail", command=self.connect_gmail).pack(side=tk.LEFT)
        ttk.Button(email_buttons, text="Create Gmail Draft", command=self.create_gmail_draft).pack(side=tk.LEFT, padx=(6, 0))
        ttk.Button(email_buttons, text="Open Email App", command=self.open_email_app).pack(side=tk.LEFT, padx=(6, 0))
        ttk.Button(email_buttons, text="Gmail Setup Help", command=self.show_gmail_setup_help).pack(side=tk.LEFT, padx=(6, 0))

        ttk.Label(fields_panel, text="PDF Fields and AI Values").pack(anchor="w")
        self.fields_tree = ttk.Treeview(fields_panel, columns=("value",), show="tree headings")
        self.fields_tree.heading("#0", text="Field")
        self.fields_tree.heading("value", text="Value")
        self.fields_tree.column("#0", width=220)
        self.fields_tree.column("value", width=320)
        self.fields_tree.pack(fill=tk.BOTH, expand=True)
        self.fields_tree.bind("<Double-1>", self.edit_selected_value)

        ttk.Label(docx_panel, text="Word Paragraphs and Last DOCX Change").pack(anchor="w")
        ttk.Label(
            docx_panel,
            text="Review the transcript before applying if exact DOCX wording matters.",
        ).pack(anchor="w", pady=(2, 6))
        self.docx_paragraph_tree = ttk.Treeview(docx_panel, columns=("preview",), show="headings", height=8)
        self.docx_paragraph_tree.heading("preview", text="Paragraph Preview")
        self.docx_paragraph_tree.column("preview", width=520)
        self.docx_paragraph_tree.pack(fill=tk.BOTH, expand=True)
        self.docx_paragraph_tree.bind("<<TreeviewSelect>>", self.on_docx_paragraph_selected)

        ttk.Label(docx_panel, text="Selected Word paragraph").pack(anchor="w", pady=(8, 0))
        self.docx_selected_paragraph_text = tk.Text(docx_panel, wrap=tk.WORD, height=5)
        self.docx_selected_paragraph_text.pack(fill=tk.BOTH, expand=True, pady=(4, 0))
        self.docx_selected_paragraph_text.configure(state=tk.DISABLED)

        ttk.Label(docx_panel, text="Last applied DOCX change").pack(anchor="w", pady=(8, 0))
        self.docx_change_summary_text = tk.Text(docx_panel, wrap=tk.WORD, height=7)
        self.docx_change_summary_text.pack(fill=tk.BOTH, expand=True, pady=(4, 0))
        self.docx_change_summary_text.configure(state=tk.DISABLED)

        preview_toolbar = ttk.Frame(preview)
        preview_toolbar.pack(fill=tk.X, pady=(0, 6))
        ttk.Button(preview_toolbar, text="Previous Page", command=self.previous_page).pack(side=tk.LEFT)
        ttk.Button(preview_toolbar, text="Next Page", command=self.next_page).pack(side=tk.LEFT, padx=(6, 0))
        ttk.Label(preview_toolbar, textvariable=self.page_status_var).pack(side=tk.LEFT, padx=(12, 0))

        preview_body = ttk.Frame(preview)
        preview_body.pack(fill=tk.BOTH, expand=True)
        self.preview_canvas = tk.Canvas(preview_body, background="#f4f4f4")
        preview_y = ttk.Scrollbar(preview_body, orient=tk.VERTICAL, command=self.preview_canvas.yview)
        preview_x = ttk.Scrollbar(preview_body, orient=tk.HORIZONTAL, command=self.preview_canvas.xview)
        self.preview_canvas.configure(yscrollcommand=preview_y.set, xscrollcommand=preview_x.set)
        self.preview_canvas.bind("<MouseWheel>", self._scroll_preview_with_mousewheel)
        self.preview_canvas.bind("<Button-4>", self._scroll_preview_with_mousewheel)
        self.preview_canvas.bind("<Button-5>", self._scroll_preview_with_mousewheel)
        self.preview_canvas.bind("<Button-1>", self.edit_preview_field)
        self.preview_canvas.grid(row=0, column=0, sticky="nsew")
        preview_y.grid(row=0, column=1, sticky="ns")
        preview_x.grid(row=1, column=0, sticky="ew")
        preview_body.columnconfigure(0, weight=1)
        preview_body.rowconfigure(0, weight=1)

        ttk.Label(main, textvariable=self.status_var).pack(anchor="w", pady=(8, 0))

    def toggle_panel(self, panel_name: str) -> None:
        if panel_name in self.collapsed_panels:
            self.collapsed_panels.remove(panel_name)
        else:
            self.collapsed_panels.add(panel_name)
        self._refresh_collapsible_panels()

    def _refresh_collapsible_panels(self) -> None:
        panel_labels = {
            "email": ("[>] Email", "[v] Email"),
            "fields": ("[>] PDF Fields", "[v] PDF Fields"),
            "docx": ("[>] DOCX Edits", "[v] DOCX Edits"),
        }

        for panel_name, (show_label, hide_label) in panel_labels.items():
            self.panel_button_vars[panel_name].set(show_label if panel_name in self.collapsed_panels else hide_label)

        for pane, _weight in self.collapsible_panes.values():
            try:
                self.left_column.forget(pane)
            except tk.TclError:
                pass

        for panel_name in ("email", "fields", "docx"):
            if panel_name in self.collapsed_panels:
                continue
            pane, weight = self.collapsible_panes[panel_name]
            self.left_column.add(pane, weight=weight)

        hidden_count = len(self.collapsed_panels)
        self.status_var.set(f"{hidden_count} panel{'s' if hidden_count != 1 else ''} collapsed")

    def _load_api_key(self) -> str:
        if self.api_key_file.exists():
            return self.api_key_file.read_text(encoding="utf-8").strip()
        return ""

    def _scroll_preview_with_mousewheel(self, event) -> str:
        self.commit_preview_editor()
        if getattr(event, "num", None) == 4 or getattr(event, "delta", 0) > 0:
            self.preview_canvas.yview_scroll(-3, "units")
        else:
            self.preview_canvas.yview_scroll(3, "units")
        return "break"

    def edit_preview_field(self, event) -> str:
        field_name = self._field_at_preview_position(event.x, event.y)
        if field_name:
            self.start_inline_preview_edit(field_name)
        else:
            self.commit_preview_editor()
        return "break"

    def _field_at_preview_position(self, canvas_x: int, canvas_y: int) -> str | None:
        page_x = self.preview_canvas.canvasx(canvas_x) / self.preview_zoom
        page_y = self.preview_canvas.canvasy(canvas_y) / self.preview_zoom

        for field_name, locations in self.pdf_field_locations.items():
            for page_index, rect in locations:
                if page_index == self.current_page_index and rect.contains(fitz.Point(page_x, page_y)):
                    return field_name
        return None

    def save_api_key(self) -> None:
        self.api_key_file.write_text(self.api_key_var.get().strip(), encoding="utf-8")
        self.status_var.set("API key saved locally")

    def test_api_key(self) -> None:
        if self.is_busy:
            return
        if not self.api_key_var.get().strip():
            messagebox.showerror("Missing API key", "Enter your Google AI Studio API key first.")
            return

        self.is_busy = True
        model_name = self.gemini_model_var.get().strip() or DEFAULT_GEMINI_MODEL
        self.status_var.set(f"Testing Gemini API key with {model_name}")
        threading.Thread(target=self._test_api_key_worker, args=(model_name,), daemon=True).start()

    def _test_api_key_worker(self, model_name: str) -> None:
        try:
            client = genai.Client(api_key=self.api_key_var.get().strip())
            response = client.models.generate_content(
                model=model_name,
                contents="Reply with only: OK",
            )
            reply = (response.text or "").strip()
            if "OK" not in reply.upper():
                raise RuntimeError(f"Gemini responded, but with unexpected text: {reply}")
            self.root.after(0, self._test_api_key_done, model_name)
        except Exception as exc:
            self.root.after(0, self._task_failed, "API key test failed", str(exc))

    def _test_api_key_done(self, model_name: str) -> None:
        self.is_busy = False
        self.status_var.set(f"API key works with {model_name}")
        messagebox.showinfo("API key works", f"Gemini responded successfully using:\n{model_name}")

    def load_pdf(self, path: str | Path | None = None) -> None:
        if path is None:
            path = filedialog.askopenfilename(title="Choose a fillable PDF", filetypes=[("PDF files", "*.pdf")])
        if not path:
            return
        path = str(path)

        try:
            reader = PdfReader(path)
            fields = reader.get_fields() or {}
        except Exception as exc:
            messagebox.showerror("PDF error", str(exc))
            return

        if not fields:
            messagebox.showerror(
                "No fillable fields found",
                "This PDF does not appear to contain fillable form fields. Use Adobe Acrobat Pro > Prepare Form first.",
            )
            return

        self.pdf_path = Path(path)
        if self.pdf_document is not None:
            self.pdf_document.close()
        self.pdf_document = fitz.open(path)
        self.pdf_fields = fields
        self.pdf_widget_values = {}
        self.pdf_field_locations = self._read_pdf_field_locations()
        for field_name in self.pdf_field_locations:
            self.pdf_fields.setdefault(field_name, {})
        self.field_values = {field_name: self._initial_pdf_field_value(field_name, details) for field_name, details in self.pdf_fields.items()}
        self.current_page_index = 0
        self.pdf_status_var.set(f"Loaded PDF: {self.pdf_path.name} ({len(fields)} fields)")
        self.refresh_fields_tree()
        self.render_current_page()

    def refresh_template_choices(self) -> None:
        templates = sorted(path.name for path in self.templates_dir.glob("*.pdf"))
        if hasattr(self, "template_combo"):
            self.template_combo.configure(values=templates)
        if templates and self.template_var.get() not in templates:
            self.template_var.set(templates[0])
        elif not templates:
            self.template_var.set("")

    def load_selected_template(self) -> None:
        template_name = self.template_var.get()
        if not template_name:
            messagebox.showerror("No template", "No saved template is selected.")
            return
        self.load_pdf(self.templates_dir / template_name)

    def save_current_as_template(self) -> None:
        if self.pdf_path is None:
            messagebox.showerror("Missing PDF", "Load or auto-prepare a PDF first.")
            return

        template_name = simpledialog.askstring("Save template", "Template file name", initialvalue=self.pdf_path.stem)
        if not template_name:
            return
        safe_name = "".join(char if char.isalnum() or char in "-_ ." else "_" for char in template_name).strip()
        if not safe_name:
            return
        if not safe_name.lower().endswith(".pdf"):
            safe_name += ".pdf"

        output_path = self.templates_dir / safe_name
        if self.pdf_path.resolve() != output_path.resolve():
            shutil.copy2(self.pdf_path, output_path)
        self.refresh_template_choices()
        self.template_var.set(output_path.name)
        self.status_var.set(f"Template saved: {output_path.name}")

    def load_docx(self, path: str | Path | None = None) -> None:
        if path is None:
            path = filedialog.askopenfilename(
                title="Choose a Word document",
                filetypes=[("DOCX files", "*.docx"), ("DOC files", "*.doc")],
            )
        if not path:
            return

        document_path = Path(path)
        if document_path.suffix.lower() == ".doc":
            messagebox.showerror(
                "Unsupported format",
                "Legacy .doc files are not supported in V6 yet. Open the file in Word and save it as .docx first.",
            )
            return

        try:
            document = Document(str(document_path))
        except Exception as exc:
            messagebox.showerror("Word document error", str(exc))
            return

        self.docx_path = document_path
        self.docx_document = document
        self.selected_docx_paragraph_index = None
        self.docx_status_var.set(f"Loaded Word document: {document_path.name} ({len(document.paragraphs)} paragraphs)")
        self.docx_selected_status_var.set("No Word paragraph selected")
        self.refresh_docx_views()
        self.status_var.set("Word document ready for voice edits")

    def reload_docx(self) -> None:
        if self.docx_path is None:
            messagebox.showerror("Missing document", "Upload a .docx file first.")
            return
        self.load_docx(self.docx_path)

    def save_docx_copy(self) -> None:
        if self.docx_document is None or self.docx_path is None:
            messagebox.showerror("Missing document", "Upload a .docx file first.")
            return

        output_path = filedialog.asksaveasfilename(
            title="Save edited DOCX",
            defaultextension=".docx",
            initialfile=f"{self.docx_path.stem}-edited.docx",
            filetypes=[("DOCX files", "*.docx")],
        )
        if not output_path:
            return

        try:
            self.docx_document.save(output_path)
            self.status_var.set(f"Edited DOCX saved: {output_path}")
            messagebox.showinfo("Save complete", f"Saved edited DOCX:\n{output_path}")
        except Exception as exc:
            messagebox.showerror("Save failed", str(exc))

    def refresh_docx_views(self) -> None:
        self.docx_paragraph_tree.delete(*self.docx_paragraph_tree.get_children())
        self._set_docx_selected_paragraph_text("")

        if self.docx_document is None:
            self.docx_selected_status_var.set("No Word paragraph selected")
            self._set_docx_change_summary({})
            return

        for index, paragraph in enumerate(self.docx_document.paragraphs, start=1):
            preview = self._docx_paragraph_preview(paragraph.text)
            self.docx_paragraph_tree.insert("", tk.END, iid=str(index), values=(f"{index}. {preview}",))

        if self.docx_document.paragraphs:
            initial_index = self.selected_docx_paragraph_index + 1 if self.selected_docx_paragraph_index is not None else 1
            if str(initial_index) in self.docx_paragraph_tree.get_children():
                self.docx_paragraph_tree.selection_set(str(initial_index))
                self.docx_paragraph_tree.focus(str(initial_index))
                self.docx_paragraph_tree.see(str(initial_index))
                self._update_docx_selected_paragraph(initial_index - 1)
            else:
                self._update_docx_selected_paragraph(0)
        else:
            self.docx_selected_status_var.set("No Word paragraph selected")

    def on_docx_paragraph_selected(self, _event=None) -> None:
        selection = self.docx_paragraph_tree.selection()
        if not selection:
            return
        self._update_docx_selected_paragraph(int(selection[0]) - 1)

    def _update_docx_selected_paragraph(self, zero_based_index: int) -> None:
        if self.docx_document is None:
            return
        if zero_based_index < 0 or zero_based_index >= len(self.docx_document.paragraphs):
            return

        self.selected_docx_paragraph_index = zero_based_index
        paragraph = self.docx_document.paragraphs[zero_based_index]
        preview = paragraph.text.strip() or "(blank paragraph)"
        self.docx_selected_status_var.set(f"Selected Word paragraph: {zero_based_index + 1}")
        self._set_docx_selected_paragraph_text(preview)

    def _set_docx_selected_paragraph_text(self, value: str) -> None:
        self.docx_selected_paragraph_text.configure(state=tk.NORMAL)
        self.docx_selected_paragraph_text.delete("1.0", tk.END)
        self.docx_selected_paragraph_text.insert("1.0", value)
        self.docx_selected_paragraph_text.configure(state=tk.DISABLED)

    def _docx_paragraph_preview(self, text: str) -> str:
        compact = re.sub(r"\s+", " ", text).strip()
        if not compact:
            return "(blank paragraph)"
        return compact if len(compact) <= 140 else compact[:137] + "..."

    def auto_prepare_flat_pdf(self) -> None:
        source = filedialog.askopenfilename(title="Choose a flat PDF to auto-prepare", filetypes=[("PDF files", "*.pdf")])
        if not source:
            return

        output_name = f"{Path(source).stem}-auto-template.pdf"
        output_path = self.templates_dir / output_name
        try:
            field_count = self._create_fillable_template_from_flat_pdf(Path(source), output_path)
        except Exception as exc:
            messagebox.showerror("Auto-prepare failed", str(exc))
            return

        if field_count == 0:
            messagebox.showwarning(
                "No fields detected",
                "No obvious blank lines or checkbox squares were detected. This form may need manual preparation in Adobe Acrobat.",
            )
            return

        self.refresh_template_choices()
        self.template_var.set(output_path.name)
        self.load_pdf(output_path)
        self.status_var.set(f"Auto-prepared template with {field_count} fields: {output_path.name}")

    def review_fields_with_gemini(self) -> None:
        if self.is_busy:
            return
        if self.pdf_path is None or not self.pdf_fields:
            messagebox.showerror("Missing PDF", "Load or auto-prepare a PDF first.")
            return
        if not self.api_key_var.get().strip():
            messagebox.showerror("Missing API key", "Enter your Google AI Studio API key first.")
            return

        self.commit_preview_editor()
        self.is_busy = True
        self.status_var.set("Sending detected fields to Gemini for review")
        threading.Thread(target=self._gemini_field_review_worker, daemon=True).start()

    def _gemini_field_review_worker(self) -> None:
        try:
            if self.pdf_path is None:
                raise RuntimeError("No PDF is loaded.")

            review_payload, image_parts = self._build_field_review_payload(self.pdf_path)
            prompt = (
                "You are reviewing a PDF form that has been auto-prepared with interactive fields. "
                "Use the page images and the detected field rectangles to improve the field layer. "
                "Return strict JSON only with this shape: "
                "{\"renames\":[{\"old_name\":\"...\",\"new_name\":\"...\",\"label\":\"...\"}],"
                "\"missing_fields\":[{\"page\":1,\"type\":\"text\",\"name\":\"...\",\"bbox\":[x0,y0,x1,y1]}],"
                "\"notes\":[\"...\"]}. "
                "Use snake_case field names. Rename vague fields like page_1_text_3 to meaningful names. "
                "Only suggest missing fields when a clear blank line, empty box, or checkbox exists. "
                "Coordinates must be PDF point coordinates matching the supplied field rectangles. "
                "Do not suggest fields for labels, titles, instructions, or filled static text.\n\n"
                f"Detected fields and page sizes:\n{json.dumps(review_payload, indent=2)}"
            )

            client = genai.Client(api_key=self.api_key_var.get().strip())
            response = client.models.generate_content(
                model=self.gemini_model_var.get().strip() or DEFAULT_GEMINI_MODEL,
                contents=[prompt, *image_parts],
                config=types.GenerateContentConfig(response_mime_type="application/json"),
            )
            review = self._parse_json_response(response.text or "")
            output_path, rename_count, missing_count = self._apply_gemini_field_review(review)
            self.root.after(0, self._gemini_field_review_done, output_path, rename_count, missing_count)
        except Exception as exc:
            self.root.after(0, self._task_failed, "Gemini field review failed", str(exc))

    def _build_field_review_payload(self, pdf_path: Path) -> tuple[dict, list[types.Part]]:
        payload_pages = []
        image_parts: list[types.Part] = []
        doc = fitz.open(pdf_path)
        try:
            for page_index, page in enumerate(doc):
                page_fields = []
                for field_name, locations in self.pdf_field_locations.items():
                    for location_page_index, rect in locations:
                        if location_page_index != page_index:
                            continue
                        field_type = "checkbox" if "checkbox" in field_name.lower() else "text"
                        page_fields.append(
                            {
                                "name": field_name,
                                "type": field_type,
                                "bbox": [round(rect.x0, 2), round(rect.y0, 2), round(rect.x1, 2), round(rect.y1, 2)],
                            }
                        )

                payload_pages.append(
                    {
                        "page": page_index + 1,
                        "width": round(page.rect.width, 2),
                        "height": round(page.rect.height, 2),
                        "text": page.get_text("text")[:4000],
                        "fields": page_fields,
                    }
                )

                pixmap = page.get_pixmap(matrix=fitz.Matrix(1, 1), alpha=False)
                image_parts.append(types.Part.from_bytes(data=pixmap.tobytes("png"), mime_type="image/png"))
        finally:
            doc.close()

        return {"pages": payload_pages}, image_parts

    def _apply_gemini_field_review(self, review: dict) -> tuple[Path, int, int]:
        if self.pdf_path is None:
            raise RuntimeError("No PDF is loaded.")

        output_path = self.templates_dir / f"{self.pdf_path.stem}-gemini-reviewed.pdf"
        rename_map = self._build_review_rename_map(review.get("renames", []))
        missing_fields = review.get("missing_fields", [])
        doc = fitz.open(self.pdf_path)
        used_names = self._collect_widget_names(doc)
        rename_count = 0
        missing_count = 0

        try:
            for page in doc:
                for widget in page.widgets() or []:
                    new_name = rename_map.get(widget.field_name)
                    if not new_name:
                        continue
                    final_name = self._unique_field_name(new_name, used_names - {widget.field_name})
                    used_names.discard(widget.field_name)
                    used_names.add(final_name)
                    widget.field_name = final_name
                    widget.update()
                    rename_count += 1

            for item in missing_fields if isinstance(missing_fields, list) else []:
                try:
                    page_number = int(item.get("page", 1))
                    page = doc[page_number - 1]
                    bbox = item.get("bbox", [])
                    if len(bbox) != 4:
                        continue
                    rect = fitz.Rect(float(bbox[0]), float(bbox[1]), float(bbox[2]), float(bbox[3]))
                    rect = rect & page.rect
                    if rect.width < 6 or rect.height < 6:
                        continue
                    field_type = str(item.get("type", "text")).lower()
                    field_name = self._unique_field_name(self._sanitize_field_name(str(item.get("name", "missing_field"))), used_names)
                    widget = fitz.Widget()
                    widget.rect = rect
                    widget.field_name = field_name
                    widget.field_value = ""
                    widget.field_type = fitz.PDF_WIDGET_TYPE_CHECKBOX if field_type == "checkbox" else fitz.PDF_WIDGET_TYPE_TEXT
                    page.add_widget(widget)
                    used_names.add(field_name)
                    missing_count += 1
                except Exception:
                    continue

            doc.save(output_path, deflate=True, garbage=4)
        finally:
            doc.close()

        return output_path, rename_count, missing_count

    def _build_review_rename_map(self, renames) -> dict[str, str]:
        rename_map: dict[str, str] = {}
        if isinstance(renames, dict):
            items = renames.items()
            for old_name, new_value in items:
                new_name = new_value.get("new_name", "") if isinstance(new_value, dict) else str(new_value)
                sanitized = self._sanitize_field_name(new_name)
                if old_name and sanitized:
                    rename_map[str(old_name)] = sanitized
            return rename_map

        if isinstance(renames, list):
            for item in renames:
                if not isinstance(item, dict):
                    continue
                old_name = str(item.get("old_name", ""))
                new_name = self._sanitize_field_name(str(item.get("new_name", "")))
                if old_name and new_name:
                    rename_map[old_name] = new_name
        return rename_map

    def _collect_widget_names(self, doc: fitz.Document) -> set[str]:
        names: set[str] = set()
        for page in doc:
            for widget in page.widgets() or []:
                if widget.field_name:
                    names.add(widget.field_name)
        return names

    def _sanitize_field_name(self, field_name: str) -> str:
        cleaned = "".join(char.lower() if char.isalnum() else "_" for char in field_name.strip())
        cleaned = "_".join(part for part in cleaned.split("_") if part)
        if cleaned and cleaned[0].isdigit():
            cleaned = f"field_{cleaned}"
        return cleaned

    def _unique_field_name(self, field_name: str, used_names: set[str]) -> str:
        base = self._sanitize_field_name(field_name) or "field"
        candidate = base
        suffix = 2
        while candidate in used_names:
            candidate = f"{base}_{suffix}"
            suffix += 1
        return candidate

    def _gemini_field_review_done(self, output_path: Path, rename_count: int, missing_count: int) -> None:
        self.is_busy = False
        self.refresh_template_choices()
        self.template_var.set(output_path.name)
        self.load_pdf(output_path)
        self.status_var.set(f"Gemini review saved: {output_path.name} ({rename_count} renamed, {missing_count} added)")

    def _create_fillable_template_from_flat_pdf(self, source_path: Path, output_path: Path) -> int:
        doc = fitz.open(source_path)
        total_fields = 0
        for page_index, page in enumerate(doc):
            candidates = self._detect_flat_form_fields(page)
            for candidate_index, (field_type, rect) in enumerate(candidates, start=1):
                widget = fitz.Widget()
                widget.rect = rect
                widget.field_name = f"page_{page_index + 1}_{field_type}_{candidate_index}"
                widget.field_value = ""
                if field_type == "checkbox":
                    widget.field_type = fitz.PDF_WIDGET_TYPE_CHECKBOX
                else:
                    widget.field_type = fitz.PDF_WIDGET_TYPE_TEXT
                page.add_widget(widget)
                total_fields += 1

        if total_fields:
            doc.save(output_path, deflate=True, garbage=4)
        doc.close()
        return total_fields

    def _detect_flat_form_fields(self, page: fitz.Page) -> list[tuple[str, fitz.Rect]]:
        candidates: list[tuple[str, fitz.Rect]] = []
        for drawing in page.get_drawings():
            for item in drawing.get("items", []):
                kind = item[0]
                if kind == "l":
                    p1, p2 = item[1], item[2]
                    if abs(p1.y - p2.y) <= 1.5 and abs(p2.x - p1.x) >= 50:
                        x0, x1 = sorted((p1.x, p2.x))
                        y = min(p1.y, page.rect.height - 4)
                        candidates.append(("text", fitz.Rect(x0, max(0, y - 16), x1, min(page.rect.height, y + 4))))
                elif kind == "re":
                    rect = fitz.Rect(item[1])
                    width = rect.width
                    height = rect.height
                    if 7 <= width <= 24 and 7 <= height <= 24 and abs(width - height) <= 6:
                        candidates.append(("checkbox", rect))
                    elif width >= 45 and 8 <= height <= 40:
                        candidates.append(("text", rect))

        return self._deduplicate_detected_fields(candidates)

    def _deduplicate_detected_fields(self, candidates: list[tuple[str, fitz.Rect]]) -> list[tuple[str, fitz.Rect]]:
        deduped: list[tuple[str, fitz.Rect]] = []
        for field_type, rect in candidates:
            if rect.width < 6 or rect.height < 6:
                continue
            if any(field_type == existing_type and rect.intersects(existing_rect) for existing_type, existing_rect in deduped):
                continue
            deduped.append((field_type, rect))
        return deduped

    def _initial_pdf_field_value(self, field_name: str, details: dict) -> str:
        widget_value = self.pdf_widget_values.get(field_name, "")
        if widget_value:
            return widget_value

        value = self._clean_pdf_value(details.get("/V", ""))
        if value:
            return value

        return ""

    def _clean_pdf_value(self, value) -> str:
        if value in (None, ""):
            return ""
        text = str(value)
        if text.startswith("/"):
            text = text[1:]
        return text

    def _read_pdf_field_locations(self) -> dict[str, list[tuple[int, fitz.Rect]]]:
        locations: dict[str, list[tuple[int, fitz.Rect]]] = {}
        if self.pdf_document is None:
            return locations

        for page_index, page in enumerate(self.pdf_document):
            for widget in page.widgets() or []:
                if not widget.field_name:
                    continue
                locations.setdefault(widget.field_name, []).append((page_index, fitz.Rect(widget.rect)))
                widget_value = self._clean_pdf_value(getattr(widget, "field_value", ""))
                if widget_value and not self.pdf_widget_values.get(widget.field_name):
                    self.pdf_widget_values[widget.field_name] = widget_value
        return locations

    def previous_page(self) -> None:
        if self.pdf_document is None or self.current_page_index <= 0:
            return
        self.commit_preview_editor()
        self.current_page_index -= 1
        self.render_current_page()

    def next_page(self) -> None:
        if self.pdf_document is None or self.current_page_index >= self.pdf_document.page_count - 1:
            return
        self.commit_preview_editor()
        self.current_page_index += 1
        self.render_current_page()

    def render_current_page(self) -> None:
        self.clear_preview_editor()
        self.preview_canvas.delete("all")
        if self.pdf_document is None:
            self.page_status_var.set("No page loaded")
            return

        page = self.pdf_document[self.current_page_index]
        matrix = fitz.Matrix(self.preview_zoom, self.preview_zoom)
        pixmap = page.get_pixmap(matrix=matrix, alpha=False)
        self.preview_image = tk.PhotoImage(data=pixmap.tobytes("png"))
        self.preview_canvas.create_image(0, 0, anchor=tk.NW, image=self.preview_image)
        self.preview_canvas.configure(scrollregion=(0, 0, pixmap.width, pixmap.height))
        self.page_status_var.set(f"Page {self.current_page_index + 1} of {self.pdf_document.page_count}")

        for field_name, locations in self.pdf_field_locations.items():
            for page_index, rect in locations:
                if page_index != self.current_page_index:
                    continue
                self._draw_field_overlay(field_name, rect)

    def _draw_field_overlay(self, field_name: str, rect: fitz.Rect) -> None:
        x0 = rect.x0 * self.preview_zoom
        y0 = rect.y0 * self.preview_zoom
        x1 = rect.x1 * self.preview_zoom
        y1 = rect.y1 * self.preview_zoom
        value = self.field_values.get(field_name, "")
        outline = "#2d6cdf" if value else "#c46a00"
        fill = "#eaf2ff" if value else "#fff6e8"
        self.preview_canvas.create_rectangle(x0, y0, x1, y1, outline=outline, width=2)
        if value:
            self.preview_canvas.create_rectangle(x0 + 1, y0 + 1, x1 - 1, y1 - 1, fill=fill, outline="")
            self.preview_canvas.create_text(
                x0 + 4,
                y0 + 3,
                anchor=tk.NW,
                text=value,
                fill="#111111",
                width=max(20, int(x1 - x0 - 8)),
            )

    def start_inline_preview_edit(self, field_name: str) -> None:
        self.commit_preview_editor()
        rect = self._field_rect_on_current_page(field_name)
        if rect is None:
            return

        x0 = rect.x0 * self.preview_zoom
        y0 = rect.y0 * self.preview_zoom
        x1 = rect.x1 * self.preview_zoom
        y1 = rect.y1 * self.preview_zoom
        editor = ttk.Entry(self.preview_canvas)
        editor.insert(0, self.field_values.get(field_name, ""))
        editor.icursor(tk.END)
        editor.bind("<Return>", lambda _event: self.commit_preview_editor())
        editor.bind("<Escape>", lambda _event: self.clear_preview_editor())
        editor.bind("<FocusOut>", lambda _event: self.commit_preview_editor())

        self.preview_editor = editor
        self.preview_editor_field = field_name
        self.preview_editor_window = self.preview_canvas.create_window(
            x0 + 2,
            y0 + 2,
            anchor=tk.NW,
            window=editor,
            width=max(24, int(x1 - x0 - 4)),
            height=max(22, int(y1 - y0 - 4)),
        )
        editor.focus_set()

    def _field_rect_on_current_page(self, field_name: str) -> fitz.Rect | None:
        for page_index, rect in self.pdf_field_locations.get(field_name, []):
            if page_index == self.current_page_index:
                return rect
        return None

    def commit_preview_editor(self) -> str:
        if self.preview_editor is None or self.preview_editor_field is None:
            return "break"

        field_name = self.preview_editor_field
        self.field_values[field_name] = self.preview_editor.get()
        self.clear_preview_editor()
        self.refresh_fields_tree()
        if self.fields_tree.exists(field_name):
            self.fields_tree.selection_set(field_name)
            self.fields_tree.see(field_name)
        self.render_current_page()
        return "break"

    def clear_preview_editor(self) -> str:
        editor = self.preview_editor
        editor_window = self.preview_editor_window
        self.preview_editor = None
        self.preview_editor_window = None
        self.preview_editor_field = None
        if editor_window is not None:
            self.preview_canvas.delete(editor_window)
        if editor is not None:
            editor.destroy()
        return "break"

    def refresh_fields_tree(self) -> None:
        self.fields_tree.delete(*self.fields_tree.get_children())
        for field_name in sorted(self.pdf_fields):
            self.fields_tree.insert("", tk.END, iid=field_name, text=field_name, values=(self.field_values.get(field_name, ""),))

    def edit_selected_value(self, _event=None) -> None:
        selected = self.fields_tree.selection()
        if not selected:
            return
        self.edit_field_value(selected[0])

    def edit_field_value(self, field_name: str) -> None:
        self.commit_preview_editor()
        current_value = self.field_values.get(field_name, "")
        new_value = simpledialog.askstring("Edit field value", field_name, initialvalue=current_value)
        if new_value is None:
            return
        self.field_values[field_name] = new_value
        self.refresh_fields_tree()
        if self.fields_tree.exists(field_name):
            self.fields_tree.selection_set(field_name)
            self.fields_tree.see(field_name)
        self.render_current_page()

    def start_recording(self) -> None:
        if self.recording or self.is_busy:
            return

        self.recorded_chunks = []

        def callback(indata, frames, time_info, status):
            if status:
                self.root.after(0, self.recording_status_var.set, f"Mic warning: {status}")
            self.recorded_chunks.append(indata.copy())

        try:
            self.record_stream = sd.InputStream(
                samplerate=self.sample_rate,
                channels=1,
                dtype="int16",
                callback=callback,
            )
            self.record_stream.start()
            self.recording = True
            self.recording_status_var.set("Recording... click Stop and Transcribe when done")
            self.status_var.set("Recording")
        except Exception as exc:
            self.recording = False
            self.record_stream = None
            messagebox.showerror("Mic error", str(exc))

    def stop_recording(self) -> None:
        if not self.recording or self.record_stream is None:
            return

        try:
            self.record_stream.stop()
            self.record_stream.close()
        finally:
            self.record_stream = None
            self.recording = False

        if not self.recorded_chunks:
            self.recording_status_var.set("Mic idle")
            messagebox.showerror("No audio", "No microphone audio was captured.")
            return

        audio = np.concatenate(self.recorded_chunks, axis=0)
        temp_wav = tempfile.NamedTemporaryFile(prefix="whisper_form_", suffix=".wav", delete=False)
        temp_wav_path = Path(temp_wav.name)
        temp_wav.close()

        with wave.open(str(temp_wav_path), "wb") as wav_file:
            wav_file.setnchannels(1)
            wav_file.setsampwidth(2)
            wav_file.setframerate(self.sample_rate)
            wav_file.writeframes(audio.tobytes())

        self.is_busy = True
        self.recording_status_var.set("Transcribing recording...")
        self.status_var.set("Loading Whisper and transcribing")
        threading.Thread(target=self._transcribe_worker, args=(temp_wav_path,), daemon=True).start()

    def _transcribe_worker(self, audio_path: Path) -> None:
        try:
            if self.whisper_model is None:
                self.whisper_model = WhisperModel(
                    self.whisper_model_var.get(),
                    device="cpu",
                    compute_type="int8",
                    download_root=str(self.model_cache_dir),
                )
            segments, _ = self.whisper_model.transcribe(str(audio_path), beam_size=5)
            transcript = "\n".join(segment.text.strip() for segment in segments if segment.text).strip()
            self.root.after(0, self._transcription_done, transcript)
        except Exception as exc:
            self.root.after(0, self._task_failed, "Transcription failed", str(exc))

    def _transcription_done(self, transcript: str) -> None:
        self.is_busy = False
        self.transcript_text.delete("1.0", tk.END)
        self.transcript_text.insert("1.0", transcript)
        self.recording_status_var.set("Mic idle")
        self.status_var.set("Transcript ready")

    def rewrite_transcript_with_gemini(self, mode: str) -> None:
        if self.is_busy:
            return
        if not self.api_key_var.get().strip():
            messagebox.showerror("Missing API key", "Enter your Google AI Studio API key first.")
            return

        transcript = self.transcript_text.get("1.0", tk.END).strip()
        if not transcript:
            messagebox.showerror("Missing transcript", "Record or type a transcript first.")
            return

        self.is_busy = True
        label = "email" if mode == "email" else "grammar cleanup"
        self.status_var.set(f"Sending transcript to Gemini for {label}")
        threading.Thread(target=self._rewrite_transcript_worker, args=(mode, transcript), daemon=True).start()

    def _rewrite_transcript_worker(self, mode: str, transcript: str) -> None:
        try:
            if mode == "email":
                instruction = (
                    "Turn the transcript into a clear, professional email draft. "
                    "Preserve the speaker's intent and important details. "
                    "Return strict JSON only with this shape: "
                    "{\"to\":\"\",\"cc\":\"\",\"bcc\":\"\",\"subject\":\"...\",\"body\":\"...\"}. "
                    "Put email addresses in to/cc/bcc only if they are clearly spoken. "
                    "Add a concise subject line, greeting, body, and closing in body. "
                    "Do not invent names, dates, attachments, commitments, or facts not present in the transcript. "
                    "If a recipient name is unknown, use a generic greeting."
                )
                result_label = "Email draft ready"
            else:
                instruction = (
                    "Clean up this voice-to-text transcript. "
                    "Fix grammar, punctuation, capitalization, and likely speech recognition mistakes. "
                    "Keep the same meaning, same facts, same order, and same overall voice. "
                    "Do not add new details, remove important details, or turn it into a different format."
                )
                result_label = "Transcript cleaned up"

            prompt = f"{instruction}\n\nTranscript:\n{transcript}"
            client = genai.Client(api_key=self.api_key_var.get().strip())
            request = {
                "model": self.gemini_model_var.get().strip() or DEFAULT_GEMINI_MODEL,
                "contents": prompt,
            }
            if mode == "email":
                request["config"] = types.GenerateContentConfig(response_mime_type="application/json")
            response = client.models.generate_content(**request)
            rewritten = (response.text or "").strip()
            if not rewritten:
                raise RuntimeError("Gemini returned an empty response.")
            if mode == "email":
                email_draft = self._parse_json_response(rewritten)
                self.root.after(0, self._email_draft_done, email_draft, result_label)
                return
            self.root.after(0, self._rewrite_transcript_done, rewritten, result_label)
        except Exception as exc:
            self.root.after(0, self._task_failed, "Gemini transcript rewrite failed", str(exc))

    def _rewrite_transcript_done(self, rewritten: str, result_label: str) -> None:
        self.is_busy = False
        self.transcript_text.delete("1.0", tk.END)
        self.transcript_text.insert("1.0", rewritten)
        self.status_var.set(result_label)

    def apply_transcript_to_docx(self) -> None:
        if self.is_busy:
            return
        if self.docx_document is None:
            messagebox.showerror("Missing document", "Upload a .docx file first.")
            return
        if not self.api_key_var.get().strip():
            messagebox.showerror("Missing API key", "Enter your Google AI Studio API key first.")
            return

        instruction = self.transcript_text.get("1.0", tk.END).strip()
        if not instruction:
            messagebox.showerror("Missing instruction", "Record or type the DOCX edit instruction first.")
            return

        self.is_busy = True
        self.status_var.set("Sending DOCX edit instruction to Gemini")
        threading.Thread(target=self._docx_edit_worker, args=(instruction,), daemon=True).start()

    def _docx_edit_worker(self, instruction: str) -> None:
        try:
            if self.docx_document is None:
                raise RuntimeError("No Word document is loaded.")

            paragraphs_payload = []
            for index, paragraph in enumerate(self.docx_document.paragraphs, start=1):
                paragraphs_payload.append({"index": index, "text": paragraph.text.strip()})

            selected_hint = self.selected_docx_paragraph_index + 1 if self.selected_docx_paragraph_index is not None else None
            prompt = (
                "You map a spoken Word document edit instruction to one precise local edit. "
                "Return strict JSON only with this shape: "
                "{\"operation\":\"...\",\"paragraph_index\":1,\"anchor_text\":\"...\",\"new_text\":\"...\",\"explanation\":\"...\"}. "
                "Allowed operation values are: insert_after_sentence, insert_before_sentence, replace_sentence, replace_paragraph, insert_paragraph_after, insert_paragraph_before, clarify. "
                "paragraph_index must be one of the provided paragraph indexes. "
                "Copy anchor_text exactly from the chosen paragraph when possible. "
                "Copy new_text from the user's instruction as literally as possible. Do not improve wording, summarize, or add facts. "
                "If the request is ambiguous or you cannot identify one safe target, return operation clarify and explain what is missing.\n\n"
                f"Selected paragraph hint: {selected_hint}\n\n"
                f"Document paragraphs:\n{json.dumps(paragraphs_payload, ensure_ascii=True)}\n\n"
                f"Instruction:\n{instruction}"
            )

            client = genai.Client(api_key=self.api_key_var.get().strip())
            response = client.models.generate_content(
                model=self.gemini_model_var.get().strip() or DEFAULT_GEMINI_MODEL,
                contents=prompt,
                config=types.GenerateContentConfig(response_mime_type="application/json"),
            )
            plan_data = self._parse_json_response(response.text or "")
            plan = self._coerce_docx_edit_plan(plan_data)
            self.root.after(0, self._docx_edit_done, plan)
        except Exception as exc:
            self.root.after(0, self._task_failed, "DOCX edit failed", str(exc))

    def _coerce_docx_edit_plan(self, payload: dict) -> DocxEditPlan:
        operation = str(payload.get("operation", "")).strip()
        if operation not in EDIT_OPERATIONS:
            raise ValueError(f"Unsupported DOCX edit operation from Gemini: {operation}")

        paragraph_index = int(payload.get("paragraph_index", 0) or 0)
        if operation != "clarify" and paragraph_index <= 0:
            raise ValueError("Gemini did not return a valid DOCX paragraph_index.")

        return DocxEditPlan(
            operation=operation,
            paragraph_index=paragraph_index,
            anchor_text=str(payload.get("anchor_text", "") or "").strip(),
            new_text=str(payload.get("new_text", "") or "").strip(),
            explanation=str(payload.get("explanation", "") or "").strip(),
        )

    def _docx_edit_done(self, plan: DocxEditPlan) -> None:
        self.is_busy = False

        if plan.operation == "clarify":
            self.status_var.set("DOCX instruction needs clarification")
            self._set_docx_change_summary(
                {
                    "operation": plan.operation,
                    "paragraph_index": plan.paragraph_index,
                    "anchor_text": plan.anchor_text,
                    "new_text": plan.new_text,
                    "explanation": plan.explanation,
                }
            )
            messagebox.showinfo("Clarify instruction", plan.explanation or "Gemini could not identify one safe DOCX edit.")
            return

        if self.docx_document is None:
            messagebox.showerror("Missing document", "The Word document is no longer loaded.")
            return

        if plan.paragraph_index > len(self.docx_document.paragraphs):
            messagebox.showerror("Edit failed", f"Paragraph {plan.paragraph_index} is outside the current Word document.")
            return

        paragraph = self.docx_document.paragraphs[plan.paragraph_index - 1]
        self._apply_docx_edit_plan_to_paragraph(paragraph, plan)
        self.selected_docx_paragraph_index = plan.paragraph_index - 1
        self.refresh_docx_views()
        self.status_var.set(f"Applied {plan.operation} to Word paragraph {plan.paragraph_index}")
        self._set_docx_change_summary(
            {
                "operation": plan.operation,
                "paragraph_index": plan.paragraph_index,
                "anchor_text": plan.anchor_text,
                "new_text": plan.new_text,
                "explanation": plan.explanation,
            }
        )

    def _apply_docx_edit_plan_to_paragraph(self, paragraph: Paragraph, plan: DocxEditPlan) -> None:
        if plan.operation == "replace_paragraph":
            self._replace_docx_paragraph_text(paragraph, plan.new_text)
            return

        if plan.operation == "insert_paragraph_after":
            self._insert_docx_paragraph_after(paragraph, plan.new_text)
            return

        if plan.operation == "insert_paragraph_before":
            self._insert_docx_paragraph_before(paragraph, plan.new_text)
            return

        if not plan.anchor_text:
            raise ValueError("Gemini did not provide anchor_text for the DOCX sentence-level edit.")

        current_text = paragraph.text
        if not current_text.strip():
            raise ValueError("The chosen Word paragraph is blank and cannot be edited at sentence level.")

        if plan.operation == "replace_sentence":
            updated = self._replace_docx_first_occurrence(current_text, plan.anchor_text, plan.new_text)
        elif plan.operation == "insert_after_sentence":
            updated = self._replace_docx_first_occurrence(
                current_text,
                plan.anchor_text,
                self._join_docx_sentences(plan.anchor_text, plan.new_text),
            )
        elif plan.operation == "insert_before_sentence":
            updated = self._replace_docx_first_occurrence(
                current_text,
                plan.anchor_text,
                self._join_docx_sentences(plan.new_text, plan.anchor_text),
            )
        else:
            raise ValueError(f"Unsupported local DOCX edit operation: {plan.operation}")

        self._replace_docx_paragraph_text(paragraph, updated)

    def _replace_docx_paragraph_text(self, paragraph: Paragraph, new_text: str) -> None:
        style = paragraph.style
        alignment = paragraph.alignment
        paragraph.text = new_text
        paragraph.style = style
        paragraph.alignment = alignment

    def _insert_docx_paragraph_after(self, paragraph: Paragraph, text: str) -> Paragraph:
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_paragraph = Paragraph(new_p, paragraph._parent)
        if paragraph.style is not None:
            new_paragraph.style = paragraph.style
        if text:
            new_paragraph.add_run(text)
        return new_paragraph

    def _insert_docx_paragraph_before(self, paragraph: Paragraph, text: str) -> Paragraph:
        new_p = OxmlElement("w:p")
        paragraph._p.addprevious(new_p)
        new_paragraph = Paragraph(new_p, paragraph._parent)
        if paragraph.style is not None:
            new_paragraph.style = paragraph.style
        if text:
            new_paragraph.add_run(text)
        return new_paragraph

    def _replace_docx_first_occurrence(self, text: str, anchor_text: str, replacement: str) -> str:
        pattern = re.compile(re.escape(anchor_text), re.IGNORECASE)
        updated, count = pattern.subn(replacement, text, count=1)
        if count != 1:
            raise ValueError("The target sentence could not be found exactly in the selected Word paragraph. Review the instruction and try again.")
        return updated

    def _join_docx_sentences(self, left: str, right: str) -> str:
        if not left:
            return right
        if not right:
            return left
        separator = ""
        if not left.endswith((" ", "\n")) and not right.startswith((" ", "\n")):
            separator = " "
        return f"{left}{separator}{right}"

    def _set_docx_change_summary(self, summary: dict) -> None:
        self.docx_change_summary_text.configure(state=tk.NORMAL)
        self.docx_change_summary_text.delete("1.0", tk.END)
        if summary:
            self.docx_change_summary_text.insert("1.0", json.dumps(summary, indent=2))
        self.docx_change_summary_text.configure(state=tk.DISABLED)

    def _email_draft_done(self, email_draft: dict, result_label: str) -> None:
        self.is_busy = False
        self.email_to_var.set(str(email_draft.get("to", "") or ""))
        self.email_cc_var.set(str(email_draft.get("cc", "") or ""))
        self.email_bcc_var.set(str(email_draft.get("bcc", "") or ""))
        self.email_subject_var.set(str(email_draft.get("subject", "") or ""))
        body = str(email_draft.get("body", "") or "")
        self.email_body_text.delete("1.0", tk.END)
        self.email_body_text.insert("1.0", body)
        self.transcript_text.delete("1.0", tk.END)
        self.transcript_text.insert("1.0", body)
        self.status_var.set(result_label)

    def connect_gmail(self) -> None:
        if self.is_busy:
            return
        self.is_busy = True
        self.status_var.set("Connecting to Gmail")
        threading.Thread(target=self._connect_gmail_worker, daemon=True).start()

    def _connect_gmail_worker(self) -> None:
        try:
            self._get_gmail_service()
            self.root.after(0, self._gmail_connected)
        except Exception as exc:
            self.root.after(0, self._task_failed, "Gmail connection failed", str(exc))

    def _gmail_connected(self) -> None:
        self.is_busy = False
        self.status_var.set("Gmail connected")
        messagebox.showinfo("Gmail connected", "Gmail is connected. You can now create drafts.")

    def create_gmail_draft(self) -> None:
        if self.is_busy:
            return
        draft = self._current_email_draft()
        if not draft["to"]:
            messagebox.showerror("Missing recipient", "Enter a recipient in the To field before creating a Gmail draft.")
            return
        if not draft["subject"] and not draft["body"]:
            messagebox.showerror("Missing email", "Create or type an email subject/body first.")
            return

        self.is_busy = True
        self.status_var.set("Creating Gmail draft")
        threading.Thread(target=self._create_gmail_draft_worker, args=(draft,), daemon=True).start()

    def _create_gmail_draft_worker(self, draft: dict[str, str]) -> None:
        try:
            service = self._get_gmail_service()
            message = EmailMessage()
            message["To"] = draft["to"]
            if draft["cc"]:
                message["Cc"] = draft["cc"]
            if draft["bcc"]:
                message["Bcc"] = draft["bcc"]
            message["Subject"] = draft["subject"]
            message.set_content(draft["body"])

            encoded = base64.urlsafe_b64encode(message.as_bytes()).decode("utf-8")
            result = service.users().drafts().create(userId="me", body={"message": {"raw": encoded}}).execute()
            self.root.after(0, self._gmail_draft_created, result.get("id", ""))
        except Exception as exc:
            self.root.after(0, self._task_failed, "Gmail draft failed", str(exc))

    def _gmail_draft_created(self, draft_id: str) -> None:
        self.is_busy = False
        self.status_var.set(f"Gmail draft created: {draft_id}")
        messagebox.showinfo("Gmail draft created", "Draft created in Gmail. Review it in Gmail before sending.")

    def open_email_app(self) -> None:
        draft = self._current_email_draft()
        if not draft["subject"] and not draft["body"]:
            messagebox.showerror("Missing email", "Create or type an email subject/body first.")
            return
        mailto = f"mailto:{quote(draft['to'])}?subject={quote(draft['subject'])}&body={quote(draft['body'])}"
        if draft["cc"]:
            mailto += f"&cc={quote(draft['cc'])}"
        if draft["bcc"]:
            mailto += f"&bcc={quote(draft['bcc'])}"
        webbrowser.open(mailto)
        self.status_var.set("Opened email draft in default email app")

    def show_gmail_setup_help(self) -> None:
        help_window = tk.Toplevel(self.root)
        help_window.title("Gmail Setup Help")
        help_window.geometry("760x620")

        container = ttk.Frame(help_window, padding=12)
        container.pack(fill=tk.BOTH, expand=True)

        text = tk.Text(container, wrap=tk.WORD)
        scrollbar = ttk.Scrollbar(container, orient=tk.VERTICAL, command=text.yview)
        text.configure(yscrollcommand=scrollbar.set)
        text.grid(row=0, column=0, sticky="nsew")
        scrollbar.grid(row=0, column=1, sticky="ns")
        container.columnconfigure(0, weight=1)
        container.rowconfigure(0, weight=1)

        instructions = f"""Gmail Draft Setup - Option 1: Use Your Own Google OAuth Client

This setup lets each person use their own Google Cloud OAuth credentials. Their Gmail login token stays on their own computer.

What you need:
- A Google account
- Internet access
- This V6 app folder:
    {self.app_dir}

Step 1 - Open Google Cloud Console
Go to:
https://console.cloud.google.com/

Step 2 - Create or select a project
- Click the project selector at the top.
- Choose an existing project or create a new one.
- A name like "Whisper Voice To Form Gmail" is fine.

Step 3 - Enable the Gmail API
- In Google Cloud Console, search for "Gmail API".
- Open Gmail API.
- Click Enable.

Step 4 - Configure the OAuth consent screen
- Go to APIs & Services > OAuth consent screen.
- Choose External unless this is a Google Workspace internal app.
- Fill in the required app name, user support email, and developer contact email.
- For a private/test setup, keep the app in Testing mode.

Step 5 - Add yourself as a test user
- On the OAuth consent screen page, find Test users.
- Add the Gmail address that will use this app.
- If someone else is using the app, add their Gmail address too.

Step 6 - Create OAuth credentials
- Go to APIs & Services > Credentials.
- Click Create Credentials.
- Choose OAuth client ID.
- Application type: Desktop app.
- Give it a name such as "Whisper Voice To Form Desktop".
- Click Create.

Step 7 - Download the JSON file
- Download the OAuth client JSON file.
- Rename it exactly to:
    gmail-credentials.json

Step 8 - Put the file in the V6 app folder
Place gmail-credentials.json here:
{self.app_dir}

Step 9 - Connect Gmail in the app
- Return to V6.
- Click Connect Gmail.
- Your browser should open.
- Log into Gmail.
- Approve Gmail compose access.

Step 10 - Create drafts
- Use Make Email to create an editable email draft.
- Review the To, Subject, and Body fields.
- Click Create Gmail Draft.
- Open Gmail and review the draft before sending.

Privacy notes:
- This app creates Gmail drafts only. It does not auto-send email.
- Your Gmail token is saved locally as .gmail-token.json.
- Do not share gmail-credentials.json or .gmail-token.json publicly.
- For legal, medical, student, or client information, review your confidentiality rules before sending text to Gemini.
"""
        text.insert("1.0", instructions)
        text.configure(state=tk.DISABLED)

        buttons = ttk.Frame(container)
        buttons.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(8, 0))
        ttk.Button(buttons, text="Open Google Cloud Console", command=lambda: webbrowser.open("https://console.cloud.google.com/")).pack(side=tk.LEFT)
        ttk.Button(buttons, text="Close", command=help_window.destroy).pack(side=tk.RIGHT)

    def show_api_key_help(self) -> None:
        help_window = tk.Toplevel(self.root)
        help_window.title("Gemini API Key Help")
        help_window.geometry("760x620")

        container = ttk.Frame(help_window, padding=12)
        container.pack(fill=tk.BOTH, expand=True)

        text = tk.Text(container, wrap=tk.WORD)
        scrollbar = ttk.Scrollbar(container, orient=tk.VERTICAL, command=text.yview)
        text.configure(yscrollcommand=scrollbar.set)
        text.grid(row=0, column=0, sticky="nsew")
        scrollbar.grid(row=0, column=1, sticky="ns")
        container.columnconfigure(0, weight=1)
        container.rowconfigure(0, weight=1)

        instructions = f"""Gemini API Key Setup - Google AI Studio

This app uses a Google AI Studio API key for Gemini requests. This is separate from Gmail OAuth.

What you need:
- A Google account
- Internet access
- This V6 app folder if you want to know where the key is saved locally:
    {self.app_dir}

Step 1 - Open Google AI Studio
Go to:
https://aistudio.google.com/

Step 2 - Sign in
- Sign in with the Google account you want to use for Gemini.

Step 3 - Open the API key page
- In Google AI Studio, click Get API key or API keys.
- If prompted, choose an existing Google Cloud project or create/select one.

Step 4 - Create the key
- Click Create API key.
- Copy the new key right away.

Step 5 - Paste it into V6
- Return to the app.
- Paste the key into the Google AI Studio API key box.
- Click Save Key.
- Click Test Key to confirm it works.

Step 6 - If the test fails
- Make sure you copied the full key.
- Try the default Gemini model first.
- If your account does not have access to a model, switch the model field to one your account can use, such as gemini-flash-latest.

Important notes:
- This Gemini API key is not the same thing as Gmail login or Gmail OAuth.
- Do not share the key or commit it to Git.
- This app stores the key locally in .gemini-api-key inside this app folder.
- Google may require project setup, billing, or region-supported access depending on the model.
"""
        text.insert("1.0", instructions)
        text.configure(state=tk.DISABLED)

        buttons = ttk.Frame(container)
        buttons.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(8, 0))
        ttk.Button(buttons, text="Open Google AI Studio", command=lambda: webbrowser.open("https://aistudio.google.com/")).pack(side=tk.LEFT)
        ttk.Button(buttons, text="Close", command=help_window.destroy).pack(side=tk.RIGHT)

    def _current_email_draft(self) -> dict[str, str]:
        return {
            "to": self.email_to_var.get().strip(),
            "cc": self.email_cc_var.get().strip(),
            "bcc": self.email_bcc_var.get().strip(),
            "subject": self.email_subject_var.get().strip(),
            "body": self.email_body_text.get("1.0", tk.END).strip(),
        }

    def _get_gmail_service(self):
        try:
            from google.auth.transport.requests import Request
            from google.oauth2.credentials import Credentials
            from google_auth_oauthlib.flow import InstalledAppFlow
            from googleapiclient.discovery import build
        except ImportError as exc:
            raise RuntimeError("Gmail dependencies are missing. Re-run install_windows_v6.bat.") from exc

        credentials = None
        if self.gmail_token_file.exists():
            credentials = Credentials.from_authorized_user_file(str(self.gmail_token_file), GMAIL_SCOPES)

        if credentials and credentials.expired and credentials.refresh_token:
            credentials.refresh(Request())

        if not credentials or not credentials.valid:
            if not self.gmail_credentials_file.exists():
                raise RuntimeError(
                    "Missing gmail-credentials.json. Create a Google Cloud OAuth Desktop client, "
                    "download its JSON file, rename it to gmail-credentials.json, and place it in the V6 app folder."
                )
            flow = InstalledAppFlow.from_client_secrets_file(str(self.gmail_credentials_file), GMAIL_SCOPES)
            credentials = flow.run_local_server(port=0)

        self.gmail_token_file.write_text(credentials.to_json(), encoding="utf-8")
        return build("gmail", "v1", credentials=credentials)

    def send_transcript_to_ai(self) -> None:
        if self.is_busy:
            return
        if not self.pdf_fields:
            messagebox.showerror("Missing PDF", "Upload a fillable PDF first.")
            return
        if not self.api_key_var.get().strip():
            messagebox.showerror("Missing API key", "Enter your Google AI Studio API key first.")
            return

        transcript = self.transcript_text.get("1.0", tk.END).strip()
        if not transcript:
            messagebox.showerror("Missing transcript", "Record or type a transcript first.")
            return

        self.is_busy = True
        self.status_var.set("Sending transcript and PDF fields to Gemini")
        threading.Thread(target=self._gemini_worker, args=(transcript,), daemon=True).start()

    def _gemini_worker(self, transcript: str) -> None:
        try:
            client = genai.Client(api_key=self.api_key_var.get().strip())
            field_payload = []
            for name, details in self.pdf_fields.items():
                field_payload.append(
                    {
                        "name": name,
                        "type": str(details.get("/FT", "")),
                        "label": str(details.get("/TU", "")),
                        "current_value": self.field_values.get(name, ""),
                    }
                )

            prompt = (
                "You fill Adobe PDF form fields from a voice transcript. "
                "Return strict JSON only. Use exact PDF field names as keys. "
                "Use empty strings for fields that are not answered. "
                "Do not invent personal information. For checkboxes, use Yes or Off.\n\n"
                f"PDF fields:\n{json.dumps(field_payload, indent=2)}\n\n"
                f"Transcript:\n{transcript}"
            )
            response = client.models.generate_content(
                model=self.gemini_model_var.get().strip() or DEFAULT_GEMINI_MODEL,
                contents=prompt,
                config=types.GenerateContentConfig(response_mime_type="application/json"),
            )
            values = self._parse_json_response(response.text or "")
            clean_values = {name: str(values.get(name, "")) for name in self.pdf_fields}
            self.root.after(0, self._gemini_done, clean_values)
        except Exception as exc:
            self.root.after(0, self._task_failed, "Gemini request failed", str(exc))

    def _parse_json_response(self, text: str) -> dict:
        stripped = text.strip()
        if stripped.startswith("```"):
            stripped = stripped.strip("`")
            if stripped.lower().startswith("json"):
                stripped = stripped[4:].strip()
        parsed = json.loads(stripped)
        if not isinstance(parsed, dict):
            raise ValueError("Gemini did not return a JSON object.")
        return parsed

    def _gemini_done(self, values: dict[str, str]) -> None:
        self.is_busy = False
        self.field_values.update(values)
        self.refresh_fields_tree()
        self.render_current_page()
        self.status_var.set("AI values ready for review. Double-click a value to edit it.")

    def export_filled_pdf(self) -> None:
        if self.pdf_path is None:
            messagebox.showerror("Missing PDF", "Upload a fillable PDF first.")
            return

        output_path = filedialog.asksaveasfilename(
            title="Save filled PDF",
            defaultextension=".pdf",
            initialfile=f"{self.pdf_path.stem}-filled.pdf",
            filetypes=[("PDF files", "*.pdf")],
        )
        if not output_path:
            return

        try:
            reader = PdfReader(str(self.pdf_path))
            writer = PdfWriter()
            for page in reader.pages:
                writer.add_page(page)

            if "/AcroForm" in reader.trailer["/Root"]:
                writer._root_object.update({NameObject("/AcroForm"): reader.trailer["/Root"]["/AcroForm"]})
                writer._root_object["/AcroForm"].update({NameObject("/NeedAppearances"): BooleanObject(True)})

            values = {name: value for name, value in self.field_values.items() if value != ""}
            for page in writer.pages:
                writer.update_page_form_field_values(page, values)

            with open(output_path, "wb") as output_file:
                writer.write(output_file)

            self.status_var.set(f"Filled PDF saved: {output_path}")
            messagebox.showinfo("Export complete", f"Saved filled PDF:\n{output_path}")
        except Exception as exc:
            messagebox.showerror("PDF export failed", str(exc))

    def _task_failed(self, title: str, error: str) -> None:
        self.is_busy = False
        self.status_var.set(title)
        self.recording_status_var.set("Mic idle")
        messagebox.showerror(title, error)


def main() -> None:
    root = tk.Tk()
    GeminiFormFillerApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
